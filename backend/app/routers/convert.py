"""
Convert Router - Image & Data Conversion API Endpoints
======================================================

This router handles image and data conversion requests.

API Flow:
1. Client sends POST request with file + target_format
2. Server validates the file type and size
3. Server converts using the appropriate service
4. If authenticated: uploads to Supabase Storage + saves history
5. Server returns the converted file as a downloadable response
"""

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Request
from fastapi.responses import Response, FileResponse
from typing import Optional, List
import os
if os.name == 'nt':
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = (
        r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    )

from app.services.image_converter import (
    convert_image,
    get_mime_type,
    validate_format,
    SUPPORTED_FORMATS,
    scrub_image_metadata,
)
from app.services.data_converter import (
    convert_data,
    validate_data_format,
    get_data_content_type,
    SUPPORTED_DATA_FORMATS,
)
from app.services.document_converter import convert_document as _convert_doc
from app.utils.file_utils import (
    get_format_from_filename,
    get_output_filename,
    get_content_type,
    validate_file_size,
    sanitize_filename,
    get_data_format_from_filename,
    get_data_output_filename,
    save_to_history,
    HISTORY_DIR,
    fetch_cloud_file
)
from app.config import settings
from app.utils.auth import get_optional_user
from app.services.storage_service import upload_file, BUCKET_ORIGINALS, BUCKET_CONVERTED
from app.services.conversion_history_service import (
    create_conversion_record,
    complete_conversion,
    fail_conversion,
    increment_user_stats,
)


router = APIRouter()


# ---- Legacy local history (kept for backwards compat / non-auth users) ----

@router.get("/history")
async def get_history():
    """List all files in history directory"""
    try:
        if not os.path.exists(HISTORY_DIR):
            return {"files": []}

        files = []
        for filename in os.listdir(HISTORY_DIR):
            file_path = os.path.join(HISTORY_DIR, filename)
            if os.path.isfile(file_path):
                stat = os.stat(file_path)
                files.append({
                    "name": filename,
                    "size": stat.st_size,
                    "created_at": stat.st_ctime,
                    "modified_at": stat.st_mtime
                })

        files.sort(key=lambda x: x["created_at"], reverse=True)
        return {"files": files}
    except Exception as e:
        return {"error": str(e), "files": []}

@router.get("/history/{filename}")
async def download_history_file(filename: str):
    """Download a specific file from history"""
    file_path = os.path.join(HISTORY_DIR, sanitize_filename(filename))
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/octet-stream'
    )


# =============================================================================
# IMAGE CONVERSION
# =============================================================================

@router.post("/image")
async def convert_image_endpoint(
    request: Request,
    target_format: str = Form(
        ...,
        description="Target format: png, jpg, jpeg, or gif"
    ),
    file: Optional[UploadFile] = File(None, description="Image file to convert"),
    cloud_url: Optional[str] = Form(None),
    filename: Optional[str] = Form(None),
    source_format: Optional[str] = Form(
        default=None,
        description="Source format (auto-detected if not provided)"
    ),
    #exif data srubbing for privacy mode NEW

    privacy_mode: bool = Form(
        default=False, 
        description="Enable Privacy Mode to strip sensitive EXIF metadata"
    ),

    width: Optional[int] = Form(default=None, description="Target width in pixels (auto-scales height if blank)"),
    height: Optional[int] = Form(default=None, description="Target height in pixels (auto-scales width if blank)"),
    quality: int = Form(default=95, ge=1, le=100, description="Compression quality (1-100, affects JPEGs)"),
):
    """
    Convert an image to a different format.

    **Supported Formats**: PNG, JPG, JPEG, GIF

    **How it works**:
    1. Upload your image file
    2. Optionally specify the source format (auto-detected from file if omitted)
    3. Specify the target format you want
    4. Receive the converted image as a download

    **Notes**:
    - PNG → JPG: Transparency is flattened to white background
    - GIF → Other: Only first frame is converted (animations not preserved)
    - Any → GIF: Colors are reduced to 256 (GIF limitation)
    """
    # --- Auth (optional) ---
    user_id = await get_optional_user(request)

    # Validate target format
    target_format = target_format.lower()
    if not validate_format(target_format):
        raise HTTPException(
            status_code=400,
            detail=f"Invalid target format '{target_format}'. Allowed: {', '.join(SUPPORTED_FORMATS)}"
        )

    # Validate source format if provided
    if source_format:
        source_format = source_format.lower()
        if not validate_format(source_format):
            raise HTTPException(
                status_code=400,
                detail=f"Invalid source format '{source_format}'. Allowed: {', '.join(SUPPORTED_FORMATS)}"
            )

    # Ensure file or cloud_url is provided
    if not file and not cloud_url:
        raise HTTPException(status_code=400, detail="Must provide either file or cloud_url")
        
    if cloud_url:
        file = fetch_cloud_file(cloud_url, filename or "cloud_image")

    # Auto-detect source format from filename if not provided
    if not source_format:
        source_format = get_format_from_filename(file.filename or "")
        if not source_format:
            raise HTTPException(
                status_code=400,
                detail="Could not detect source format. Please provide source_format parameter."
            )

    # Read file contents
    try:
        file_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read file: {str(e)}")

    # --- 🛡️ PRIVACY MODE INTERCEPT --- meta data scrubber NEW
    if privacy_mode:
        try:
            file_bytes = scrub_image_metadata(file_bytes)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to scrub EXIF data: {str(e)}")

    # Validate file size (Done AFTER scrub in case scrubbing reduced size)
    if not validate_file_size(len(file_bytes), settings.MAX_FILE_SIZE):
        max_mb = settings.MAX_FILE_SIZE / (1024 * 1024)
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size is {max_mb:.0f}MB."
        )

    # Check if source and target are the same
    is_same_format = source_format == target_format or (source_format == "jpeg" and target_format == "jpg") or (source_format == "jpg" and target_format == "jpeg")
    
    # Check if the user is actually asking the API to do work (tweaking)
    is_tweaking = privacy_mode or (width is not None) or (height is not None) or (quality != 95)
    
    if is_same_format and not is_tweaking:
        raise HTTPException(
            status_code=400,
            detail=f"Source and target formats are the same ({source_format}). No conversion needed unless you are applying privacy mode, resizing, or compressing."
        )
    # --- Supabase: create pending record + upload original ---
    conversion_id = None
    if user_id:
        conversion_id = create_conversion_record(
            user_id=user_id,
            original_filename=file.filename or "unknown",
            original_format=source_format,
            converted_format=target_format,
            file_size_original=len(file_bytes),
        )
        # Upload original file to storage
        original_path = upload_file(BUCKET_ORIGINALS, file_bytes, user_id, file.filename or "original")
        if conversion_id and original_path:
            from app.services.supabase_service import get_supabase
            get_supabase().table("conversions").update(
                {"original_file_url": original_path}
            ).eq("id", conversion_id).execute()

    # Perform conversion
    try:
        converted_bytes, detected_format = convert_image(
            file_bytes=file_bytes,
            source_format=source_format,
            target_format=target_format,
            width=width,     # <-- ADD THIS
            height=height,   # <-- ADD THIS
            quality=quality  # <-- ADD THIS
        )
    except ValueError as e:
        if conversion_id:
            fail_conversion(conversion_id, str(e))
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        if conversion_id:
            fail_conversion(conversion_id, str(e))
        raise HTTPException(
            status_code=500,
            detail=f"Conversion failed: {str(e)}"
        )

    # Generate output filename
    output_filename = get_output_filename(file.filename or "converted", target_format)
    output_filename = sanitize_filename(output_filename)

    # --- Supabase: upload converted + complete record ---
    if user_id and conversion_id:
        converted_path = upload_file(BUCKET_CONVERTED, converted_bytes, user_id, output_filename)
        if converted_path:
            complete_conversion(conversion_id, converted_path, len(converted_bytes))
            increment_user_stats(user_id, len(converted_bytes))
    else:
        # Fallback: save to local history for non-auth users
        save_to_history(converted_bytes, output_filename)

    # Return the converted image
    return Response(
        content=converted_bytes,
        media_type=get_content_type(target_format),
        headers={
            "Content-Disposition": f'attachment; filename="{output_filename}"'
        }
    )


# =============================================================================
# DATA CONVERSION
# =============================================================================

@router.post("/data")
async def convert_data_endpoint(
    request: Request,
    target_format: str = Form(
        ...,
        description="Target format: json, csv, xlsx, or xml"
    ),
    file: Optional[UploadFile] = File(None, description="Data file to convert"),
    cloud_url: Optional[str] = Form(None),
    filename: Optional[str] = Form(None),
    source_format: Optional[str] = Form(
        default=None,
        description="Source format (auto-detected if not provided)"
    ),
):
    """
    Convert between data formats using hub-and-spoke architecture.

    **Supported Formats**: JSON, CSV, XLSX, XML

    **Architecture**: Input → DataFrame (Hub) → Output
    """
    # --- Auth (optional) ---
    user_id = await get_optional_user(request)

    # Validate target format
    target_format = target_format.lower()
    if not validate_data_format(target_format):
        raise HTTPException(
            status_code=400,
            detail=f"Invalid target format '{target_format}'. Allowed: {', '.join(SUPPORTED_DATA_FORMATS)}"
        )

    # Validate source format if provided
    if source_format:
        source_format = source_format.lower()
        if not validate_data_format(source_format):
            raise HTTPException(
                status_code=400,
                detail=f"Invalid source format '{source_format}'. Allowed: {', '.join(SUPPORTED_DATA_FORMATS)}"
            )

    # Auto-detect source format from filename if not provided
    if not source_format:
        source_format = get_data_format_from_filename(file.filename or "")
        if not source_format:
            raise HTTPException(
                status_code=400,
                detail="Could not detect source format. Please provide source_format parameter."
            )
    # Ensure file or cloud_url is provided
    if not file and not cloud_url:
        raise HTTPException(status_code=400, detail="Must provide either file or cloud_url")

    if cloud_url:
        file = fetch_cloud_file(cloud_url, filename or "cloud_data")

    # Read file contents
    try:
        file_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read file: {str(e)}")

    # Validate file size
    if not validate_file_size(len(file_bytes), settings.MAX_FILE_SIZE):
        max_mb = settings.MAX_FILE_SIZE / (1024 * 1024)
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size is {max_mb:.0f}MB."
        )

    # Check if source and target are the same
    if source_format == target_format:
        raise HTTPException(
            status_code=400,
            detail=f"Source and target formats are the same ({source_format}). No conversion needed."
        )

    # --- Supabase: create pending record + upload original ---
    conversion_id = None
    if user_id:
        conversion_id = create_conversion_record(
            user_id=user_id,
            original_filename=file.filename or "unknown",
            original_format=source_format,
            converted_format=target_format,
            file_size_original=len(file_bytes),
        )
        original_path = upload_file(BUCKET_ORIGINALS, file_bytes, user_id, file.filename or "original")
        if conversion_id and original_path:
            from app.services.supabase_service import get_supabase
            get_supabase().table("conversions").update(
                {"original_file_url": original_path}
            ).eq("id", conversion_id).execute()

    # Perform conversion
    try:
        converted_bytes, row_count, col_count = convert_data(
            file_bytes=file_bytes,
            source_format=source_format,
            target_format=target_format,
        )
    except ValueError as e:
        if conversion_id:
            fail_conversion(conversion_id, str(e))
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        if conversion_id:
            fail_conversion(conversion_id, str(e))
        raise HTTPException(
            status_code=500,
            detail=f"Conversion failed: {str(e)}"
        )

    # Generate output filename
    output_filename = get_data_output_filename(file.filename or "converted", target_format)
    output_filename = sanitize_filename(output_filename)

    # --- Supabase: upload converted + complete record ---
    if user_id and conversion_id:
        converted_path = upload_file(BUCKET_CONVERTED, converted_bytes, user_id, output_filename)
        if converted_path:
            complete_conversion(conversion_id, converted_path, len(converted_bytes))
            increment_user_stats(user_id, len(converted_bytes))
    else:
        save_to_history(converted_bytes, output_filename)

    # Return the converted data
    return Response(
        content=converted_bytes,
        media_type=get_data_content_type(target_format),
        headers={
            "Content-Disposition": f'attachment; filename="{output_filename}"',
            "X-Row-Count": str(row_count),
            "X-Column-Count": str(col_count),
        }
    )


# =============================================================================
# OCR CONVERSION
# =============================================================================

from fastapi.responses import StreamingResponse
import io
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document

@router.post("/ocr")
async def convert_ocr_endpoint(
    request: Request,
    file: UploadFile = File(..., description="File to OCR"),
    format: Optional[str] = 'docx'
):
    """
    Extract editable text from scanned PDFs or images using OCR.
    """
    user_id = await get_optional_user(request)

    try:
        file_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read file: {str(e)}")

    if not validate_file_size(len(file_bytes), settings.MAX_FILE_SIZE):
        max_mb = settings.MAX_FILE_SIZE / (1024 * 1024)
        raise HTTPException(status_code=413, detail=f"File too large. Maximum size is {max_mb:.0f}MB.")

    extracted_texts = []
    filename_lower = (file.filename or "").lower()
    is_pdf = filename_lower.endswith('.pdf')

    try:
        if is_pdf:
            pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page_num in range(len(pdf_doc)):
                page = pdf_doc.load_page(page_num)
                text = page.get_text().strip()
                if len(text) >= 30:
                    extracted_texts.append(text)
                else:
                    pix = page.get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img).strip()
                    extracted_texts.append(ocr_text)
            pdf_doc.close()
        else:
            img = Image.open(io.BytesIO(file_bytes))
            ocr_text = pytesseract.image_to_string(img).strip()
            extracted_texts.append(ocr_text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OCR processing failed: {str(e)}")

    full_text = "\n\n".join(extracted_texts)

    if format == 'txt':
        return Response(content=full_text, media_type="text/plain")

    # Default to docx
    doc = Document()
    doc.add_heading("OCR Extracted Text", level=1)
    for p_text in extracted_texts:
        if p_text:
            doc.add_paragraph(p_text)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    output_filename = sanitize_filename(get_output_filename(file.filename or "ocr", "docx"))

    return StreamingResponse(
        doc_io,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{output_filename}"'}
    )



@router.get("/formats")
async def get_supported_formats():
    """Get list of supported image formats."""
    return {
        "supported_formats": list(SUPPORTED_FORMATS),
        "notes": {
            "jpg_jpeg": "JPG and JPEG are equivalent (both output as .jpg)",
            "transparency": "PNG supports transparency, JPG/JPEG do not",
            "animation": "GIF animation is not preserved when converting to other formats",
            "colors": "GIF is limited to 256 colors"
        }
    }


@router.get("/data/formats")
async def get_data_formats():
    """Get list of supported data formats and conversion info."""
    return {
        "supported_formats": list(SUPPORTED_DATA_FORMATS),
        "architecture": "hub-and-spoke (Input → DataFrame → Output)",
        "notes": {
            "json": "Supports array of objects or single object, nested objects are flattened",
            "csv": "Comma-separated values, first row is headers",
            "xlsx": "Excel format, only first sheet is read/written",
            "xml": "Expected structure: <root><row><col>value</col></row></root>"
        }
    }
# REMOTE FETCH ENDPOINTS
# =============================================================================

@router.post("/remote-fetch")
async def remote_fetch_convert(
    url: str = Form(..., description="URL of the file to fetch and convert"),
    target_format: Optional[str] = Form(None, description="Target format (optional - auto-detected if not provided)"),
):
    """
    Fetch a file from a URL and convert it to a different format.
    
    This endpoint:
    1. Downloads the file from the provided URL
    2. Auto-detects the file type
    3. Converts it to the target format (or suggests options if not specified)
    4. Returns the converted file
    
    **Supported URLs**: Any publicly accessible HTTP/HTTPS URL
    **Supported formats**: Images (PNG, JPG, JPEG, GIF), Documents (PDF, DOCX), Data (JSON, CSV, XLSX, XML)
    
    **Examples**:
    - Google Drive: https://drive.google.com/file/d/FILE_ID/view
    - Dropbox: https://www.dropbox.com/s/FILE_ID/filename.pdf
    - Direct links: https://example.com/image.png
    
    **Notes**:
    - Large files may take time to download
    - Some cloud services require special sharing settings
    - Google Drive links need to be set to "Anyone with the link can view"
    """
    if not url or not url.startswith(('http://', 'https://')):
        raise HTTPException(status_code=400, detail="Invalid URL. Must be a valid HTTP/HTTPS URL.")

    try:
        # Fetch the file from URL
        # For Google Drive URLs, the last path segment is "view" (not a real filename),
        # so we must rely on content-based format detection after downloading.
        if 'drive.google.com' in url:
            filename = "remote_file"
        else:
            filename = url.split('/')[-1] or "remote_file"
            if '?' in filename:
                filename = filename.split('?')[0]
            if not filename or '.' not in filename:
                filename = "remote_file"

        file = fetch_cloud_file(url, filename)

        # Read file contents
        file_bytes = await file.read()

        # Validate file size
        if not validate_file_size(len(file_bytes), settings.MAX_FILE_SIZE):
            max_mb = settings.MAX_FILE_SIZE / (1024 * 1024)
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size is {max_mb:.0f}MB."
            )

        # Auto-detect file type from content and filename
        detected_format = None
        content_type = None

        # Check filename extension first
        filename_lower = file.filename.lower()

        # Image formats
        if filename_lower.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp')):
            detected_format = get_format_from_filename(file.filename)
            if filename_lower.endswith('.webp'):
                detected_format = "webp"
            content_type = "image"
        # Data formats
        elif filename_lower.endswith(('.json', '.csv', '.xlsx', '.xml')):
            detected_format = get_data_format_from_filename(file.filename)
            content_type = "data"
        # Document formats (basic detection)
        elif filename_lower.endswith(('.pdf', '.docx')):
            if filename_lower.endswith('.pdf'):
                detected_format = "pdf"
                content_type = "document"
            elif filename_lower.endswith('.docx'):
                detected_format = "docx"
                content_type = "document"

        # If we couldn't detect from filename, try content analysis
        if not detected_format:
            # Check for common file signatures
            if file_bytes.startswith(b'\x89PNG'):
                detected_format = "png"
                content_type = "image"
            elif file_bytes.startswith(b'\xff\xd8'):
                detected_format = "jpg"
                content_type = "image"
            elif file_bytes.startswith(b'GIF8'):
                detected_format = "gif"
                content_type = "image"
            elif len(file_bytes) >= 12 and file_bytes.startswith(b'RIFF') and file_bytes[8:12] == b'WEBP':
                detected_format = "webp"
                content_type = "image"
            elif file_bytes.startswith(b'%PDF'):
                detected_format = "pdf"
                content_type = "document"
            elif file_bytes.startswith(b'PK\x03\x04'):  # ZIP signature (DOCX/XLSX)
                if b'word/' in file_bytes[:1000]:  # DOCX
                    detected_format = "docx"
                    content_type = "document"
                elif b'xl/' in file_bytes[:1000]:  # XLSX
                    detected_format = "xlsx"
                    content_type = "data"
            elif file_bytes.strip().startswith(b'{') or file_bytes.strip().startswith(b'['):
                detected_format = "json"
                content_type = "data"
            elif b',' in file_bytes[:100]:  # Simple CSV detection
                detected_format = "csv"
                content_type = "data"

        if not detected_format:
            raise HTTPException(
                status_code=400,
                detail="Could not detect file type. Supported formats: PNG, JPG, GIF, PDF, DOCX, JSON, CSV, XLSX, XML"
            )

        # If no target format specified, suggest conversions
        if not target_format:
            suggestions = []
            if content_type == "image":
                suggestions = ["png", "jpg", "gif"]
                suggestions = [fmt for fmt in suggestions if fmt != detected_format]
            elif content_type == "document":
                if detected_format == "pdf":
                    suggestions = ["docx"]
                elif detected_format == "docx":
                    suggestions = ["pdf"]
            elif content_type == "data":
                suggestions = ["json", "csv", "xlsx", "xml"]
                suggestions = [fmt for fmt in suggestions if fmt != detected_format]

            return {
                "detected_format": detected_format,
                "content_type": content_type,
                "file_size": len(file_bytes),
                "suggested_conversions": suggestions,
                "message": f"File detected as {detected_format.upper()}. Choose a target format to convert."
            }

        # Perform conversion based on content type
        converted_bytes = None
        output_filename = None

        if content_type == "image":
            target_format = target_format.lower()

            # Image → PDF is a document conversion (handled by document_converter).
            if target_format == "pdf":
                content_type = "document"
                try:
                    output_path = await _convert_doc(
                        file_bytes,
                        file.filename or f"remote.{detected_format}",
                        detected_format,
                        target_format,
                    )
                    with open(output_path, "rb") as f:
                        converted_bytes = f.read()
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    output_filename = get_output_filename(file.filename or f"remote.{detected_format}", target_format)
                except ValueError as ve:
                    raise HTTPException(status_code=400, detail=str(ve))
                except Exception as ex:
                    raise HTTPException(status_code=400, detail=f"Unsupported conversion or failure: {str(ex)}")

            else:
                if not validate_format(target_format):
                    raise HTTPException(
                        status_code=400,
                        detail=f"Invalid target format '{target_format}'. Allowed: png, jpg, jpeg, gif, pdf"
                    )
                if detected_format == target_format:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Source and target formats are the same ({detected_format}). No conversion needed."
                    )

                converted_bytes, _ = convert_image(
                    file_bytes=file_bytes,
                    source_format=detected_format,
                    target_format=target_format,
                )
                output_filename = get_output_filename(file.filename, target_format)

        elif content_type == "data":
            target_format = target_format.lower()
            if not validate_data_format(target_format):
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid target format '{target_format}'. Allowed: json, csv, xlsx, xml"
                )

            if detected_format == target_format:
                raise HTTPException(
                    status_code=400,
                    detail=f"Source and target formats are the same ({detected_format}). No conversion needed."
                )

            converted_bytes, _, _ = convert_data(
                file_bytes=file_bytes,
                source_format=detected_format,
                target_format=target_format,
            )
            output_filename = get_data_output_filename(file.filename, target_format)
        elif content_type == "document":
            try:
                # _convert_doc now returns an absolute path to a temp file
                output_path = await _convert_doc(file_bytes, file.filename or f"remote.{detected_format}", detected_format, target_format)
                
                with open(output_path, "rb") as f:
                    converted_bytes = f.read()
                
                if os.path.exists(output_path):
                    os.remove(output_path)
                
                output_filename = get_output_filename(file.filename or f"remote.{detected_format}", target_format)
            except ValueError as ve:
                raise HTTPException(status_code=400, detail=str(ve))
            except Exception as ex:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported conversion or failure: {str(ex)}"
                )

        if not converted_bytes:
            raise HTTPException(status_code=500, detail="Conversion failed")

        # Sanitize filename for HTTP headers
        output_filename = sanitize_filename(output_filename)

        # Save to history
        save_to_history(converted_bytes, output_filename)

        # Determine content type for response
        if content_type == "image":
            media_type = get_content_type(target_format)
        elif content_type == "data":
            media_type = get_data_content_type(target_format)
        else:
            media_type = "application/octet-stream"

        # Return the converted file
        return Response(
            content=converted_bytes,
            media_type=media_type,
            headers={
                "Content-Disposition": f'attachment; filename="{output_filename}"'
            }
        )

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Remote fetch failed: {str(e)}"
        )
